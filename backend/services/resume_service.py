# backend/services/resume_service.py
"""
ATS Intelligence Engine — upgraded resume analysis service.
Uses sentence-transformers for semantic JD matching.
"""

import os
import re
from typing import List, Optional, Dict, Any, Tuple

import fitz          # PyMuPDF
import pdfplumber
import docx
from fastapi import HTTPException
from pydantic import BaseModel

# ── spaCy (optional — graceful degradation) ──────────────────────────────────
try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
except Exception:
    nlp = None

# ── sentence-transformers (optional — graceful degradation) ───────────────────
try:
    from sentence_transformers import SentenceTransformer, util as st_util
    _st_model = SentenceTransformer("all-MiniLM-L6-v2")
except Exception:
    _st_model = None

# ── textstat (optional) ───────────────────────────────────────────────────────
try:
    import textstat as _textstat
except Exception:
    _textstat = None


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SKILL VOCABULARY  (300+ curated tech terms)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SKILL_VOCAB: Dict[str, str] = {
    # Languages
    "python": "language", "java": "language", "javascript": "language",
    "typescript": "language", "c++": "language", "c#": "language",
    "go": "language", "rust": "language", "kotlin": "language",
    "swift": "language", "ruby": "language", "php": "language",
    "scala": "language", "r": "language", "matlab": "language",
    "dart": "language", "elixir": "language", "haskell": "language",
    "perl": "language", "bash": "language", "shell": "language",
    "sql": "language", "html": "language", "css": "language",
    # Frameworks / libraries
    "react": "framework", "vue": "framework", "angular": "framework",
    "next.js": "framework", "nuxt": "framework", "svelte": "framework",
    "django": "framework", "flask": "framework", "fastapi": "framework",
    "spring": "framework", "spring boot": "framework", "express": "framework",
    "node.js": "framework", "rails": "framework", "laravel": "framework",
    "tensorflow": "framework", "pytorch": "framework", "keras": "framework",
    "scikit-learn": "framework", "pandas": "framework", "numpy": "framework",
    "opencv": "framework", "hugging face": "framework", "langchain": "framework",
    "flutter": "framework", "react native": "framework",
    # Cloud & DevOps
    "aws": "cloud", "azure": "cloud", "gcp": "cloud", "google cloud": "cloud",
    "docker": "devops", "kubernetes": "devops", "terraform": "devops",
    "ansible": "devops", "jenkins": "devops", "github actions": "devops",
    "ci/cd": "devops", "linux": "devops", "nginx": "devops",
    "helm": "devops", "prometheus": "devops", "grafana": "devops",
    "datadog": "devops", "elk": "devops",
    # Databases
    "postgresql": "database", "mysql": "database", "mongodb": "database",
    "redis": "database", "elasticsearch": "database", "cassandra": "database",
    "dynamodb": "database", "sqlite": "database", "oracle": "database",
    "neo4j": "database", "influxdb": "database",
    # Tools & practices
    "git": "tool", "jira": "tool", "confluence": "tool", "figma": "tool",
    "postman": "tool", "graphql": "tool", "rest": "tool", "grpc": "tool",
    "kafka": "tool", "rabbitmq": "tool", "celery": "tool",
    "agile": "practice", "scrum": "practice", "tdd": "practice",
    "microservices": "practice", "devops": "practice", "mlops": "practice",
    "system design": "practice", "design patterns": "practice",
    # AI / ML concepts
    "machine learning": "ai", "deep learning": "ai", "nlp": "ai",
    "computer vision": "ai", "reinforcement learning": "ai",
    "llm": "ai", "generative ai": "ai", "transformers": "ai",
    "bert": "ai", "gpt": "ai", "rag": "ai", "vector database": "ai",
    "data science": "ai", "data engineering": "ai",
}

# Impact action verbs
_IMPACT_VERBS = {
    "built", "developed", "designed", "implemented", "created", "launched",
    "deployed", "optimized", "improved", "reduced", "increased", "automated",
    "delivered", "led", "managed", "architected", "integrated", "migrated",
    "refactored", "scaled", "streamlined", "achieved", "generated", "saved",
    "eliminated", "accelerated", "enhanced", "established", "collaborated",
}

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# FILE VALIDATION & TEXT EXTRACTION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def validate_file(file) -> None:
    allowed_mime = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }
    content_type = file.content_type or ""
    filename = file.filename or ""
    ext = os.path.splitext(filename)[1].lower()
    if content_type not in allowed_mime and ext not in {".pdf", ".doc", ".docx"}:
        raise HTTPException(status_code=400, detail="Unsupported file type. Upload a PDF or DOCX.")
    if file.size and file.size > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File exceeds 10 MB limit.")


def extract_text(file_path: str) -> str:
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    text = ""

    if ext == ".pdf":
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    pt = page.extract_text()
                    if pt:
                        text += pt + "\n"
        except Exception:
            pass
        if not text.strip():
            try:
                doc = fitz.open(file_path)
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            except Exception as e:
                raise HTTPException(status_code=422, detail=f"Cannot extract PDF text: {e}")
    elif ext in {".doc", ".docx"}:
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"Cannot read DOCX: {e}")
    else:
        raise HTTPException(status_code=400, detail="Unsupported file extension.")

    if not text.strip():
        raise HTTPException(status_code=422, detail="Resume appears empty or unreadable.")

    return text


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# PARSING
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _extract_contact(text: str) -> Dict[str, Optional[str]]:
    email_match = re.search(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", text)
    phone_match = re.search(
        r"(\+?\d{1,3}[\s\-]?)?(\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{4})", text
    )
    linkedin_match = re.search(r"linkedin\.com/in/[\w\-]+", text, re.IGNORECASE)
    github_match   = re.search(r"github\.com/[\w\-]+", text, re.IGNORECASE)
    return {
        "email":    email_match.group(0) if email_match else None,
        "phone":    phone_match.group(0).strip() if phone_match else None,
        "linkedin": linkedin_match.group(0) if linkedin_match else None,
        "github":   github_match.group(0) if github_match else None,
    }


def _extract_name(text: str) -> Optional[str]:
    if nlp:
        doc = nlp(text[:600])
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                return ent.text.strip()
    for line in text.splitlines():
        stripped = line.strip()
        if stripped and len(stripped.split()) <= 5 and not re.search(r"[@|{/\\]", stripped):
            return stripped
    return None


# Section heading patterns (order matters — more specific first)
_SECTION_MAP: List[Tuple[str, str]] = [
    (r"\b(professional\s+)?summary\b|\bobjective\b|\bprofile\b", "summary"),
    (r"\bskills?\b|\btechnical\s+skills?\b|\bcore\s+competenc", "skills"),
    (r"\bwork\s+(experience|history)\b|\bexperience\b|\bemployment\b|\bprofessional\s+experience\b", "experience"),
    (r"\beducation\b|\bacademic\b|\bqualification\b", "education"),
    (r"\bprojects?\b|\bpersonal\s+projects?\b|\bkey\s+projects?\b", "projects"),
    (r"\bcertif(icate|ication)s?\b|\blicens(e|es|ure)s?\b|\bachievements?\b", "certifications"),
    (r"\bpublications?\b|\bresearch\b", "publications"),
    (r"\bvolunteer\b|\bextra.?curricular\b|\bactivities\b", "activities"),
    (r"\bawards?\b|\bhonors?\b", "awards"),
]


def parse_resume(text: str) -> Dict[str, Any]:
    sections: Dict[str, List[str]] = {
        k: [] for _, k in _SECTION_MAP
    }

    name    = _extract_name(text)
    contact = _extract_contact(text)

    current_section = None
    lines = text.splitlines()

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        lower = stripped.lower()
        matched_section = None

        # Only treat short lines as headings
        if len(stripped) < 80:
            for pattern, sec in _SECTION_MAP:
                if re.search(pattern, lower):
                    matched_section = sec
                    break

        if matched_section:
            current_section = matched_section
            continue

        if current_section:
            sections[current_section].append(stripped)

    result: Dict[str, Any] = {
        "name":           name,
        "email":          contact.get("email"),
        "phone":          contact.get("phone"),
        "linkedin":       contact.get("linkedin"),
        "github":         contact.get("github"),
        "summary":        " ".join(sections["summary"]),
        "skills":         "; ".join(sections["skills"]),
        "education":      "; ".join(sections["education"]),
        "experience":     "; ".join(sections["experience"]),
        "projects":       "; ".join(sections["projects"]),
        "certifications": "; ".join(sections["certifications"]),
        "publications":   "; ".join(sections["publications"]),
        "activities":     "; ".join(sections["activities"]),
    }
    return result


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SKILL EXTRACTION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def extract_skills(text: str) -> List[Dict[str, str]]:
    """
    Match known tech skills in resume text.
    Returns list of {name, category} dicts, deduplicated and sorted.
    """
    lower_text = text.lower()
    found: Dict[str, str] = {}

    # Sort by length descending so multi-word skills match before sub-words
    for skill, category in sorted(SKILL_VOCAB.items(), key=lambda x: -len(x[0])):
        pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, lower_text):
            found[skill] = category

    return [{"name": s, "category": c} for s, c in sorted(found.items())]


def extract_jd_skills(jd_text: str) -> List[str]:
    """Extract skill names present in a job description."""
    lower_jd = jd_text.lower()
    found = []
    for skill in sorted(SKILL_VOCAB.keys(), key=lambda x: -len(x)):
        pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, lower_jd):
            found.append(skill)
    return found


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# IMPACT SCORING
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _impact_score(parsed: Dict[str, Any]) -> Tuple[float, List[str]]:
    """
    Score quantification richness in experience & projects sections.
    Returns (score 0-100, list of detected impact phrases).
    """
    target_text = " ".join([
        parsed.get("experience", "") or "",
        parsed.get("projects", "") or "",
        parsed.get("summary", "") or "",
    ])

    if not target_text.strip():
        return 0.0, []

    lines = [l.strip() for l in target_text.split(";") if l.strip()]
    if not lines:
        return 0.0, []

    impact_lines = []
    verb_count   = 0
    metric_count = 0

    for line in lines:
        low = line.lower()
        has_metric = bool(re.search(r"\d+\s*(%|x|k|m|b|ms|s|hrs?|days?|\$|users?|req|rps)", low))
        has_verb   = any(v in low.split() for v in _IMPACT_VERBS)

        if has_metric:
            metric_count += 1
            impact_lines.append(line[:120])
        if has_verb:
            verb_count += 1

    total = len(lines)
    verb_ratio   = min(verb_count / total, 1.0)   if total else 0
    metric_ratio = min(metric_count / total, 1.0) if total else 0

    score = round((verb_ratio * 50) + (metric_ratio * 50), 2)
    return score, impact_lines[:6]


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SEMANTIC JD MATCHING
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def semantic_jd_score(resume_text: str, jd_text: str) -> float:
    """
    Cosine similarity between resume and JD sentence embeddings.
    Returns score 0-100.  Falls back to keyword overlap if model unavailable.
    """
    if _st_model is not None and jd_text.strip():
        try:
            emb_r = _st_model.encode(resume_text[:4096], convert_to_tensor=True)
            emb_j = _st_model.encode(jd_text[:4096],    convert_to_tensor=True)
            sim   = float(st_util.cos_sim(emb_r, emb_j)[0][0])
            # cos_sim is -1..1; clamp and scale to 0-100
            return round(max(0.0, min(sim * 100, 100.0)), 2)
        except Exception:
            pass

    # Fallback: keyword overlap
    if jd_text.strip():
        jd_words = set(re.findall(r"\b\w{3,}\b", jd_text.lower()))
        res_words = set(re.findall(r"\b\w{3,}\b", resume_text.lower()))
        stop = {"the", "and", "for", "with", "this", "that", "are", "you", "will"}
        jd_words -= stop
        if jd_words:
            return round(len(jd_words & res_words) / len(jd_words) * 100, 2)
    return 50.0  # neutral when no JD


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# SCORING
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _keyword_match(parsed: Dict[str, Any], jd_text: Optional[str]) -> float:
    if not jd_text:
        # Score based on how many vocabulary skills appear in resume
        resume_text = " ".join(str(v) for v in parsed.values() if v)
        skills = extract_skills(resume_text)
        skill_score = min(len(skills) / 20 * 100, 100)
        return round(skill_score, 2)

    jd_words = set(re.findall(r"\b\w+\b", jd_text.lower()))
    resume_text = " ".join(str(v) for v in parsed.values() if v)
    resume_words = set(re.findall(r"\b\w+\b", resume_text.lower()))
    stop = {"the", "a", "an", "and", "or", "is", "in", "to", "of", "for",
            "with", "at", "by", "on", "as", "be", "this", "that", "we", "you"}
    jd_words -= stop
    if not jd_words:
        return 50.0
    return round(len(jd_words & resume_words) / len(jd_words) * 100, 2)


def _readability_score(text: str) -> float:
    if _textstat:
        try:
            score = _textstat.flesch_reading_ease(text)
            return round(max(0.0, min(100.0, score)), 2)
        except Exception:
            pass
    return 55.0


def _completeness_score(parsed: Dict[str, Any]) -> float:
    weights = {
        "name": 1.0, "email": 1.0, "phone": 0.8,
        "linkedin": 0.6, "github": 0.5,
        "summary": 0.8, "skills": 1.0, "education": 1.0,
        "experience": 1.5, "projects": 1.2, "certifications": 0.6,
    }
    total = sum(weights.values())
    earned = sum(w for k, w in weights.items() if parsed.get(k))
    return round(earned / total * 100, 2)


def _section_scores(parsed: Dict[str, Any]) -> Dict[str, float]:
    """Per-section quality score based on content density."""
    def density(text: str) -> float:
        if not text:
            return 0.0
        words = len(text.split())
        # Penalise very short sections, reward richness up to a cap
        return round(min(words / 60 * 100, 100), 2)

    return {
        "summary":        density(parsed.get("summary", "")),
        "skills":         density(parsed.get("skills", "")),
        "experience":     density(parsed.get("experience", "")),
        "education":      density(parsed.get("education", "")),
        "projects":       density(parsed.get("projects", "")),
        "certifications": density(parsed.get("certifications", "")),
    }


def score_resume(
    parsed: Dict[str, Any],
    jd_text: Optional[str] = None,
    resume_full_text: str = "",
    formatting_issues: Optional[List[str]] = None,
) -> Dict[str, Any]:
    formatting_issues = formatting_issues or []

    keyword_pct  = _keyword_match(parsed, jd_text)
    resume_text  = resume_full_text or " ".join(str(v) for v in parsed.values() if v)
    readability  = _readability_score(resume_text)
    completeness = _completeness_score(parsed)
    impact, impact_lines = _impact_score(parsed)
    sem_score    = semantic_jd_score(resume_text, jd_text) if jd_text else 0.0

    # Formatting (penalise per issue, max penalty 40pts)
    fmt_penalty  = min(len(formatting_issues) * 10, 40)
    formatting   = max(0, 100 - fmt_penalty)

    # Weighted ATS score
    if jd_text:
        ats_score = (
            keyword_pct  * 0.25 +
            sem_score    * 0.25 +
            readability  * 0.15 +
            completeness * 0.20 +
            formatting   * 0.10 +
            impact       * 0.05
        )
    else:
        ats_score = (
            keyword_pct  * 0.20 +
            readability  * 0.20 +
            completeness * 0.35 +
            formatting   * 0.15 +
            impact       * 0.10
        )

    return {
        "ats_score":      round(min(ats_score, 100.0), 2),
        "keyword_match":  keyword_pct,
        "semantic_score": sem_score,
        "readability":    readability,
        "completeness":   completeness,
        "impact_score":   impact,
        "formatting":     formatting,
        "impact_lines":   impact_lines,
    }


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# FORMATTING DETECTION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def detect_formatting_issues(file_path: str) -> List[str]:
    issues = []
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == ".pdf":
        try:
            doc = fitz.open(file_path)
            img_reported = col_reported = False
            for page in doc:
                blocks = page.get_text("dict").get("blocks", [])
                for b in blocks:
                    if b.get("type") == 1 and not img_reported:
                        issues.append("Images detected — ATS cannot read embedded images")
                        img_reported = True
                    if b.get("type") == 0 and b.get("lines") and not col_reported:
                        xs = [line["bbox"][0] for line in b["lines"]]
                        if xs and (max(xs) - min(xs)) > 300:
                            issues.append("Multi-column layout may confuse ATS parsers")
                            col_reported = True
            # Check for excessive pages
            if len(doc) > 2:
                issues.append(f"Resume is {len(doc)} pages — ATS prefers 1–2 pages")
            doc.close()
        except Exception:
            pass

        # Check for special characters / encoding issues
        try:
            with pdfplumber.open(file_path) as pdf:
                sample = ""
                for page in pdf.pages[:2]:
                    t = page.extract_text() or ""
                    sample += t
                non_ascii = len([c for c in sample if ord(c) > 127])
                if non_ascii > len(sample) * 0.05:
                    issues.append("High proportion of special characters may cause ATS parsing errors")
        except Exception:
            pass

    elif ext in {".doc", ".docx"}:
        try:
            doc = docx.Document(file_path)
            if doc.tables:
                issues.append("Tables detected — ATS often misreads table content")
            from docx.oxml.ns import qn
            if doc.element.body.findall(".//" + qn("w:txbxContent")):
                issues.append("Text boxes detected — ATS cannot read text box content")
            # Check header/footer for contact info
            for section in doc.sections:
                hdr = section.header
                if hdr and any(para.text.strip() for para in hdr.paragraphs):
                    issues.append("Contact info in header — some ATS skip headers")
        except Exception:
            pass

    return issues


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# FEEDBACK GENERATION
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def generate_feedback(
    score_data: Dict[str, Any],
    parsed: Optional[Dict[str, Any]] = None,
    jd_text: Optional[str] = None,
) -> Dict[str, Any]:
    strengths: List[str] = []
    weaknesses: List[str] = []
    tips: List[str] = []

    ats          = score_data.get("ats_score", 0)
    keyword      = score_data.get("keyword_match", 0)
    semantic     = score_data.get("semantic_score", 0)
    readability  = score_data.get("readability", 0)
    completeness = score_data.get("completeness", 0)
    formatting   = score_data.get("formatting", 100)
    impact       = score_data.get("impact_score", 0)

    # ATS Score
    if ats >= 80:
        strengths.append(f"Excellent ATS compatibility ({ats:.0f}/100) — top 15% of resumes")
    elif ats >= 65:
        strengths.append(f"Good ATS score ({ats:.0f}/100) — solid baseline with room to improve")
    else:
        weaknesses.append(f"ATS score of {ats:.0f}/100 — likely filtered before human review")
        tips.append("Tailor your resume to each job posting to raise your ATS score above 70.")

    # Semantic Match
    if jd_text:
        if semantic >= 75:
            strengths.append(f"Strong semantic alignment with the job description ({semantic:.0f}%)")
        elif semantic >= 55:
            weaknesses.append(f"Moderate semantic match ({semantic:.0f}%) — mirror the JD's language more closely")
            tips.append("Use the exact phrasing from the job description in your experience bullets.")
        else:
            weaknesses.append(f"Low semantic match ({semantic:.0f}%) — resume reads differently from the JD")
            tips.append("Rewrite your summary to echo the role's key responsibilities.")

    # Keywords
    if keyword >= 70:
        strengths.append("Strong keyword alignment with role requirements")
    elif keyword >= 45:
        weaknesses.append("Moderate keyword coverage — include more role-specific technical terms")
        tips.append("Add a dedicated 'Technical Skills' section listing all relevant tools.")
    else:
        weaknesses.append("Low keyword match — major ATS filters will likely reject this resume")
        tips.append("Pull 10–15 keywords directly from the job description and weave them in naturally.")

    # Impact
    if impact >= 60:
        strengths.append("Excellent use of quantified achievements — results-oriented resume")
    elif impact >= 30:
        weaknesses.append("Add more metrics to your bullets (%, $, users, time saved, scale)")
        tips.append("For each experience bullet, ask: 'How many? How much? How fast?'")
    else:
        weaknesses.append("No quantified achievements found — bullets read as duties, not results")
        tips.append("Transform duty-based bullets: 'Built API' → 'Built REST API handling 10K req/s'")

    # Readability
    if readability >= 60:
        strengths.append("Clear, professional writing style that reads well")
    else:
        weaknesses.append("Readability is below average — use shorter, punchy sentences")
        tips.append("Aim for Flesch score > 60: shorter sentences, active voice, no jargon walls.")

    # Completeness
    if completeness >= 85:
        strengths.append("All key resume sections are present and complete")
    elif completeness >= 60:
        weaknesses.append("Some sections are thin — flesh out your summary, skills, or projects")
        tips.append("Add a 3–4 sentence professional summary at the top of your resume.")
    else:
        weaknesses.append("Multiple critical sections missing — ATS will downrank significantly")
        tips.append("Ensure you have: Contact info, Summary, Skills, Experience, Education, Projects.")

    # Formatting
    if formatting >= 90:
        strengths.append("Clean, ATS-friendly single-column formatting")
    elif formatting >= 70:
        weaknesses.append("Minor formatting issues — review for ATS compatibility")
    else:
        weaknesses.append("Significant formatting issues detected — ATS may garble your content")
        tips.append("Use a plain single-column template. Avoid tables, text boxes, and columns.")

    # Parsed-specific tips
    if parsed:
        if not parsed.get("linkedin"):
            tips.append("Add your LinkedIn URL — many ATS systems score profile completeness.")
        if not parsed.get("github") and parsed.get("skills"):
            tips.append("Include your GitHub profile to back up technical skill claims.")
        if not parsed.get("summary"):
            tips.append("A professional summary helps ATS keyword scanners from the very first line.")

    return {
        "strengths":  strengths[:5],
        "weaknesses": weaknesses[:5],
        "ats_tips":   tips[:7],
    }


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# JD MATCHING  (legacy endpoint — kept for backward compat)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def match_jd(parsed: Dict[str, Any], jd_text: str) -> Dict[str, Any]:
    jd_words     = set(re.findall(r"\b\w{4,}\b", jd_text.lower()))
    resume_text  = " ".join(str(v) for v in parsed.values() if v).lower()
    resume_words = set(re.findall(r"\b\w{4,}\b", resume_text))

    stop = {"that", "this", "with", "from", "have", "will", "your", "they",
            "been", "their", "also", "more", "some", "such", "when", "than"}
    jd_words -= stop

    matched    = jd_words & resume_words
    missing    = list(jd_words - resume_words)[:20]
    match_pct  = round(len(matched) / len(jd_words) * 100, 2) if jd_words else 0.0
    sem_score  = semantic_jd_score(resume_text, jd_text)

    suggestions = []
    if missing:
        suggestions.append(f"Add these missing keywords: {', '.join(sorted(missing)[:10])}")
    if match_pct < 50:
        suggestions.append("Rewrite your experience section to mirror the job description language")
    if match_pct < 70:
        suggestions.append("Add a tailored Skills section listing technologies in the job posting")

    return {
        "match_percentage":  max(match_pct, sem_score * 0.5),
        "semantic_score":    sem_score,
        "missing_keywords":  sorted(missing),
        "suggestions":       suggestions,
    }


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# RESPONSE MODELS
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class ResumeAnalysisResponse(BaseModel):
    # Core scores (backward-compatible)
    ats_score:       float
    keyword_match:   float
    readability:     float
    completeness:    float
    formatting:      int

    # New intelligence fields
    semantic_score:  float = 0.0
    impact_score:    float = 0.0
    section_scores:  Dict[str, float] = {}
    extracted_skills: List[Dict[str, str]] = []
    missing_skills:   List[str] = []
    impact_lines:     List[str] = []
    ats_tips:         List[str] = []

    # Feedback
    strengths:          List[str]
    weaknesses:         List[str]
    parsed:             Dict[str, Any]
    formatting_issues:  List[str]


class JDMatchResponse(BaseModel):
    match_percentage:  float
    semantic_score:    float = 0.0
    missing_keywords:  List[str]
    suggestions:       List[str]


class SuggestionsResponse(BaseModel):
    optimization_tips:   List[str]
    action_verbs:        List[str]
    bullet_rewrites:     List[str]
    project_improvements: List[str]
    summary_suggestions: List[str]
